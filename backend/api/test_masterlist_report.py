"""The VPSEA masterlist renders into the office's own Word template."""
from io import BytesIO
from unittest import mock

import docx
from django.test import Client, TestCase

from api import doc_convert, masterlist_report
from api.test_support import pdf_text
from api.models import (
    AffirmativeStaffApplication, Application, Scholarship, StudentProfile,
    SystemSettings, User,
)


def _docx_text(data):
    """Every paragraph and cell of a .docx, joined — for checking what was sent
    to the converter is the real document."""
    document = docx.Document(BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    parts += [c.text for t in document.tables for row in t.rows for c in row.cells]
    return '\n'.join(parts)


class MasterlistFixtures:
    def setUp(self):
        SystemSettings.objects.create(pk=1, academic_year='26-1', active_semester='1st Semester')
        for stype, name in (
            ('Academic', 'Academic Scholarship'), ('DOST', 'DOST Scholarship'),
            ('TDP', 'TDP Scholarship'), ('CHED', 'CHED Full Merit'),
            ('Sports', 'Sports Scholarship'),
        ):
            Scholarship.objects.create(
                name=name, type=stype, category='application',
                description='x', eligibility='x', requirements=[],
            )
        self.vpsea = User.objects.create_user(
            username='vpsea@bipsu.edu.ph', email='vpsea@bipsu.edu.ph',
            password='pw', role='vpsea',
        )
        self.c = Client()
        self.assertTrue(self.c.login(email='vpsea@bipsu.edu.ph', password='pw'))

    def _scholar(self, stype, last, first, gender, sid, gwa=1.20, middle='Santos'):
        u = User.objects.create_user(
            username=f'{sid}@bipsu.edu.ph', email=f'{sid}@bipsu.edu.ph', password='pw',
            first_name=first, last_name=last, role='student',
        )
        p = StudentProfile.objects.create(
            user=u, student_id=sid, course='BSCS', year_level=3, gender=gender,
            gwa=gwa, middle_name=middle,
            barangay='Poblacion', municipality='Naval', province='Biliran',
        )
        return Application.objects.create(
            student=p, scholarship=Scholarship.objects.get(type=stype),
            status='Approved',
            award_number=f'AW-{sid}', congress_district='Lone District',
        )

    def _staff(self, full_name, gender='F', sid='EMP-01'):
        return AffirmativeStaffApplication.objects.create(
            full_name=full_name, email=f'{sid}@bipsu.edu.ph',
            contact_number='09171234567', date_of_birth='1990-01-01',
            gender=gender, course='BSIT', year_level=2, student_id=sid,
            qualified_for='Staff', status='Approved', is_nsu_staff=True,
        )

    # ── Context shape ───────────────────────────────────────────────────────


class MasterlistReportTest(MasterlistFixtures, TestCase):
    """Context building and the rendered document."""

    def test_context_supplies_every_slot_the_template_loops_over(self):
        context, _ = masterlist_report.build_context()
        for slot in masterlist_report.ALL_SLOTS:
            self.assertIn(slot, context, f'{slot} missing — Jinja would raise on its loop')
            entry = context[slot]
            self.assertIn('name', entry)
            for key in ('female', 'male', 'students'):
                self.assertIsInstance(entry[key], list)

    def test_scholars_are_split_by_gender_into_the_right_slot(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        self._scholar('Academic', 'Bautista', 'Ben', 'M', '2024-0002')
        self._scholar('DOST', 'Reyes', 'Rita', 'Female', '2024-0003')

        context, summary = masterlist_report.build_context()
        acad = context['program1']
        self.assertEqual(acad['name'], 'ACADEMIC')
        self.assertEqual([r['last_name'] for r in acad['female']], ['Cruz'])
        self.assertEqual([r['last_name'] for r in acad['male']], ['Bautista'])
        # Each gender column is numbered from 1 independently, as the form does.
        self.assertEqual(acad['female'][0]['no'], 1)
        self.assertEqual(acad['male'][0]['no'], 1)

        dost = context['program6']
        self.assertEqual(dost['name'], 'DOST')
        self.assertEqual([r['last_name'] for r in dost['female']], ['Reyes'])

        by_slot = {s['slot']: s for s in summary}
        self.assertEqual(by_slot['program1']['total'], 2)
        self.assertEqual(by_slot['program6']['total'], 1)

    def test_row_fields_match_the_template_placeholders(self):
        self._scholar('DOST', 'Cruz', 'Ana', 'F', '2024-0001')
        context, _ = masterlist_report.build_context()
        row = context['program6']['female'][0]
        self.assertEqual(row['last_name'], 'Cruz')
        self.assertEqual(row['first_name'], 'Ana')
        self.assertEqual(row['middle_name'], 'Santos')
        self.assertEqual(row['m_i'], 'S.')
        self.assertEqual(row['sex'], 'F')
        self.assertEqual(row['brgy_st'], 'Poblacion')
        self.assertEqual(row['mun'], 'Naval')
        self.assertEqual(row['municipality'], 'Naval')
        self.assertEqual(row['prov'], 'Biliran')
        self.assertEqual(row['award_number'], 'AW-2024-0001')
        self.assertEqual(row['cong_dist'], 'Lone District')
        self.assertEqual(row['number'], '2024-0001')
        self.assertEqual(row['yr'], 3)

    def test_academic_scholars_get_their_scholar_class(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001', gwa=1.20)
        self._scholar('Academic', 'Lim', 'Lea', 'F', '2024-0004', gwa=1.45)
        context, _ = masterlist_report.build_context()
        classes = [r['percent'] for r in context['program1']['female']]
        self.assertIn('Univ. Scholar', classes)
        self.assertIn('College Scholar', classes)

    def test_nsu_staff_uses_the_ungendered_students_table(self):
        self._staff('Maria Santos Lim')
        context, _ = masterlist_report.build_context()
        staff = context['program2']
        self.assertEqual(staff['name'], 'BiPSU STAFF')
        self.assertEqual(len(staff['students']), 1)
        self.assertEqual(staff['female'], [])
        row = staff['students'][0]
        self.assertEqual(row['last_name'], 'Lim')
        self.assertEqual(row['first_name'], 'Maria')
        self.assertEqual(row['m_i'], 'S.')
        self.assertEqual(row['percent'], '100%')
        self.assertEqual(row['number'], 'EMP-01')

    # ── The rendered document ───────────────────────────────────────────────
    def test_download_returns_a_word_file_with_the_scholars_in_it(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        self._scholar('TDP', 'Bautista', 'Ben', 'M', '2024-0002')
        self._staff('Maria Santos Lim')

        r = self.c.get('/vpsea/reports/download/')
        self.assertEqual(r.status_code, 200)
        self.assertIn('wordprocessingml', r['Content-Type'])
        self.assertIn('BiPSU_List_of_Scholars', r['Content-Disposition'])

        d = docx.Document(BytesIO(r.content))
        text = '\n'.join(p.text for p in d.paragraphs)
        cells = '\n'.join(c.text for t in d.tables for row in t.rows for c in row.cells)

        # The office's own headings survive verbatim.
        self.assertIn('Republic of the Philippines', text)
        self.assertIn('BILIRAN PROVINCE STATE UNIVERSITY', text)
        self.assertIn('ACADEMIC', text)
        self.assertIn('BiPSU STAFF', text)
        # No unrendered Jinja is left behind.
        self.assertNotIn('{{', text + cells)
        self.assertNotIn('{%', text + cells)
        # The scholars landed in the tables.
        self.assertIn('Cruz', cells)
        self.assertIn('Bautista', cells)
        self.assertIn('Lim', cells)

    def test_the_document_is_stamped_with_the_active_school_year(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        r = self.c.get('/vpsea/reports/download/')
        d = docx.Document(BytesIO(r.content))
        text = '\n'.join(p.text for p in d.paragraphs)
        self.assertIn('SY: 2026-2027', text)
        self.assertNotIn('2019-2020', text)

    def test_second_semester_is_reflected_in_the_headings(self):
        SystemSettings.objects.filter(pk=1).update(academic_year='26-2')
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        r = self.c.get('/vpsea/reports/download/')
        d = docx.Document(BytesIO(r.content))
        text = '\n'.join(p.text for p in d.paragraphs)
        self.assertIn('SY: 2026-2027', text)
        self.assertNotIn('1st SEM', text)
        self.assertNotIn('1ST SEM', text)

    def test_table_structure_is_the_templates_own(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        r = self.c.get('/vpsea/reports/download/')
        d = docx.Document(BytesIO(r.content))
        # Two tables per gendered program, one for the ungendered staff table.
        gendered = sum(1 for x in masterlist_report.PROGRAM_SLOTS if x[3] == 'gendered')
        ungendered = sum(1 for x in masterlist_report.PROGRAM_SLOTS if x[3] == 'students')
        self.assertEqual(len(d.tables), gendered * 2 + ungendered)
        header = [c.text.strip() for c in d.tables[0].rows[1].cells]
        self.assertEqual(header[:5], ['NO.', 'LAST NAME', 'FIRST NAME', 'MIDDLE NAME', 'SEX'])

    def test_renders_with_no_scholars_at_all(self):
        r = self.c.get('/vpsea/reports/download/')
        self.assertEqual(r.status_code, 200)
        d = docx.Document(BytesIO(r.content))
        cells = '\n'.join(c.text for t in d.tables for row in t.rows for c in row.cells)
        self.assertNotIn('{%', cells)
        self.assertNotIn('{{', cells)

    def test_only_vpsea_can_download_it(self):
        other = Client()
        User.objects.create_user(
            username='u@bipsu.edu.ph', email='u@bipsu.edu.ph', password='pw', role='unifast')
        other.login(email='u@bipsu.edu.ph', password='pw')
        r = other.get('/vpsea/reports/download/')
        self.assertEqual(r.status_code, 302)
        self.assertIn('/login/', r['Location'])


class MasterlistPreviewTest(MasterlistFixtures, TestCase):
    """The on-screen preview is driven by the same context as the document."""

    def test_preview_lists_the_programs_in_document_order(self):
        r = self.c.get('/vpsea/reports/')
        self.assertEqual(r.status_code, 200)
        headings = [s[0] for s in r.context['sections']]
        expected = [h for _slot, h, _k, _l, _hs in
                    __import__('api.masterlist_report', fromlist=['x']).PROGRAM_SLOTS]
        self.assertEqual(headings, expected)

    def test_preview_shows_the_same_rows_the_document_gets(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        self._scholar('DOST', 'Bautista', 'Ben', 'M', '2024-0002')

        r = self.c.get('/vpsea/reports/')
        self.assertEqual(r.context['grand_total'], 2)

        doc_context, _ = masterlist_report.build_context()
        # Preview rows are cells in header order, so read them by column position.
        preview = {sec[0]: sec for sec in r.context['sections']}

        def surnames(section, which):
            headers = section[1]
            column = headers.index('LAST NAME')
            return [row[column] for row in section[which]]

        self.assertEqual(surnames(preview['ACADEMIC'], 2),
                         [x['last_name'] for x in doc_context['program1']['female']])
        self.assertEqual(surnames(preview['DOST'], 3),
                         [x['last_name'] for x in doc_context['program6']['male']])

    def test_staff_section_is_not_gender_split_in_the_preview(self):
        self._staff('Maria Santos Lim')
        r = self.c.get('/vpsea/reports/')
        staff = [s for s in r.context['sections'] if s[0] == 'BiPSU STAFF'][0]
        self.assertFalse(staff[4], 'staff table has no FEMALE/MALE split')
        self.assertEqual(len(staff[2]), 1)
        self.assertEqual(staff[3], [])

    def test_preview_reports_a_missing_template(self):
        r = self.c.get('/vpsea/reports/')
        self.assertTrue(r.context['template_available'])


class MasterlistPreviewPageTest(MasterlistFixtures, TestCase):
    """The Reports tab frames the document as a PDF instead of retyping it.

    Whether the frame gets the converted Word file or the stand-in layout
    depends on the machine, so every assertion about page content pins the view
    to one path rather than reading whatever this developer happens to have
    installed.
    """

    def test_the_tab_frames_the_preview_rather_than_a_table(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        r = self.c.get('/vpsea/reports/')
        body = r.content.decode()
        self.assertIn('src="/vpsea/reports/preview/"', body)
        self.assertIn('report-preview-frame', body)
        # The old editable replica is gone: nothing on this page is typed into.
        self.assertNotIn('contenteditable', body)

    def test_the_frame_serves_the_converted_word_document(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        with mock.patch('api.doc_convert.available', return_value=True), \
                mock.patch('api.doc_convert.to_pdf',
                           return_value=b'%PDF-1.7 converted') as convert:
            r = self.c.get('/vpsea/reports/preview/')

        self.assertEqual(r.content, b'%PDF-1.7 converted')
        source, suffix = convert.call_args[0]
        self.assertEqual(suffix, '.docx')
        # What went to the converter is the generated document itself.
        self.assertTrue(source.startswith(b'PK'))
        self.assertIn('Cruz', _docx_text(source))

    def test_the_stand_in_is_used_when_no_converter_is_installed(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        self._scholar('DOST', 'Bautista', 'Ben', 'M', '2024-0002')

        with mock.patch('api.doc_convert.available', return_value=False):
            r = self.c.get('/vpsea/reports/preview/')

        self.assertEqual(r.status_code, 200)
        self.assertEqual(r['Content-Type'], 'application/pdf')
        self.assertTrue(r['Content-Disposition'].startswith('inline;'))
        self.assertTrue(r.content.startswith(b'%PDF'))

        text = pdf_text(r.content)
        self.assertIn('Cruz', text)
        self.assertIn('Bautista', text)
        self.assertIn('LIST OF SCHOLARS FOR', text)
        self.assertIn('BILIRAN PROVINCE STATE UNIVERSITY', text)
        # And it says on its face that it is not the office template.
        self.assertIn('STAND-IN LAYOUT', text)

    def test_a_failed_conversion_falls_back_instead_of_erroring(self):
        with mock.patch('api.doc_convert.available', return_value=True), \
                mock.patch('api.doc_convert.to_pdf',
                           side_effect=doc_convert.ConversionFailed('no pdf')):
            r = self.c.get('/vpsea/reports/preview/')

        self.assertEqual(r.status_code, 200)
        self.assertIn('STAND-IN LAYOUT', pdf_text(r.content))

    def test_the_preview_renders_with_no_scholars_at_all(self):
        with mock.patch('api.doc_convert.available', return_value=False):
            r = self.c.get('/vpsea/reports/preview/')
        self.assertEqual(r.status_code, 200)
        self.assertTrue(r.content.startswith(b'%PDF'))

    def test_the_preview_may_be_framed_by_our_own_pages(self):
        with mock.patch('api.doc_convert.available', return_value=False):
            r = self.c.get('/vpsea/reports/preview/')
        self.assertNotEqual(r.headers.get('X-Frame-Options'), 'DENY')

    def test_the_preview_is_closed_to_other_roles(self):
        other = Client()
        User.objects.create_user(
            username='u@bipsu.edu.ph', email='u@bipsu.edu.ph', password='pw',
            role='unifast')
        other.login(email='u@bipsu.edu.ph', password='pw')
        r = other.get('/vpsea/reports/preview/')
        self.assertEqual(r.status_code, 302)
        self.assertIn('/login/', r['Location'])


class MasterlistUnusedSlotsTest(MasterlistFixtures, TestCase):
    """The template's spare program blocks must not reach the reader."""

    def test_no_unnamed_headings_or_orphan_tables(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        r = self.c.get('/vpsea/reports/download/')
        d = docx.Document(BytesIO(r.content))

        paragraphs = [p.text.strip() for p in d.paragraphs]
        headings = [t for i, t in enumerate(paragraphs)
                    if i + 1 < len(paragraphs)
                    and paragraphs[i + 1].upper().startswith('SCHOLARSHIP GRANT')]

        # One heading per filled slot, every one of them named.
        self.assertEqual(len(headings), len(masterlist_report.PROGRAM_SLOTS))
        for heading in headings:
            stripped = heading.replace('(@)', '').replace('(*)', '').strip()
            self.assertTrue(stripped, f'unnamed program heading: {heading!r}')

        # 2 tables per gendered program, 1 for the ungendered staff table.
        gendered = sum(1 for s in masterlist_report.PROGRAM_SLOTS if s[3] == 'gendered')
        ungendered = sum(1 for s in masterlist_report.PROGRAM_SLOTS if s[3] == 'students')
        self.assertEqual(len(d.tables), gendered * 2 + ungendered)

    def test_the_internal_marker_never_reaches_the_document(self):
        r = self.c.get('/vpsea/reports/download/')
        d = docx.Document(BytesIO(r.content))
        text = '\n'.join(p.text for p in d.paragraphs)
        cells = '\n'.join(c.text for t in d.tables for row in t.rows for c in row.cells)
        self.assertNotIn(masterlist_report.UNUSED_MARKER, text + cells)
        self.assertNotIn('UNUSED', text + cells)

    def test_the_letterhead_and_page_setup_survive_the_deletions(self):
        r = self.c.get('/vpsea/reports/download/')
        d = docx.Document(BytesIO(r.content))
        text = '\n'.join(p.text for p in d.paragraphs)
        self.assertIn('Republic of the Philippines', text)
        self.assertIn('BILIRAN PROVINCE STATE UNIVERSITY', text)
        self.assertIn('Naval, Biliran', text)
        self.assertIn('LIST OF SCHOLARS FOR', text)
        # sectPr (margins, orientation) is untouched by the block removal.
        self.assertTrue(d.sections)
        self.assertIsNotNone(d.sections[0].page_width)


class MasterlistColumnShapeTest(MasterlistFixtures, TestCase):
    """Each programme's columns come from its own table in the office template."""

    def test_headers_are_read_from_the_template_not_hardcoded(self):
        headings = masterlist_report.slot_headers()
        # Every filled slot has an entry, and the shapes genuinely differ.
        for slot, _h, _k, _l, _hs in masterlist_report.PROGRAM_SLOTS:
            self.assertIn(slot, headings, slot)
        self.assertGreater(len({len(v) for v in headings.values()}), 1,
                           'all slots came out the same width — headers look hardcoded')

    def test_programmes_without_an_award_number_do_not_get_that_column(self):
        _, summary = masterlist_report.build_context()
        by_heading = {e['heading']: e['headers'] for e in summary}

        # GSIS and Sports have no award number in the office form.
        self.assertNotIn('AWARD NUMBER', by_heading['GSIS'])
        self.assertNotIn('AWARD NUMBER', by_heading['SPORTS'])
        self.assertNotIn('AWARD NUMBER', by_heading['ACADEMIC'])
        # DOST and TES do.
        self.assertIn('AWARD NUMBER', by_heading['DOST'])
        self.assertIn('AWARD NUMBER', by_heading['TES'])

    def test_the_preview_row_width_always_equals_its_header_count(self):
        self._scholar('Academic', 'Cruz', 'Ana', 'F', '2024-0001')
        self._scholar('DOST', 'Bautista', 'Ben', 'M', '2024-0002')
        self._scholar('Sports', 'Lim', 'Lea', 'F', '2024-0003')
        self._staff('Maria Santos Lim')

        r = self.c.get('/vpsea/reports/')
        self.assertEqual(r.status_code, 200)
        for title, headers, female, male, _gendered, _key in r.context['sections']:
            for row in list(female) + list(male):
                self.assertEqual(
                    len(row), len(headers),
                    f'{title}: row has {len(row)} cells for {len(headers)} headers')

    def test_a_sports_row_carries_municipality_without_a_barangay_column(self):
        self._scholar('Sports', 'Lim', 'Lea', 'F', '2024-0004')
        r = self.c.get('/vpsea/reports/')
        sports = [s for s in r.context['sections'] if s[0] == 'SPORTS'][0]
        headers, female = sports[1], sports[2]
        self.assertNotIn('BRGY./ST.', headers)
        self.assertIn('MUNICIPALITY', headers)
        self.assertEqual(female[0][headers.index('MUNICIPALITY')], 'Naval')


class StaffNameSplitTest(TestCase):
    """AffirmativeStaffApplication stores one full_name; the archives show parts."""

    def _app(self, full_name):
        return AffirmativeStaffApplication(
            full_name=full_name, email='x@bipsu.edu.ph', contact_number='09',
            date_of_birth='1990-01-01', gender='F', course='BSIT', year_level=1,
            qualified_for='Staff', status='Approved', is_nsu_staff=True,
        )

    def test_three_part_name_splits_into_last_first_middle(self):
        a = self._app('Maria Santos Lim')
        self.assertEqual(a.last_name, 'Lim')
        self.assertEqual(a.first_name, 'Maria')
        self.assertEqual(a.middle_name, 'Santos')
        self.assertEqual(a.middle_initial, 'S.')

    def test_two_part_name_has_no_middle(self):
        a = self._app('Nena Villanueva')
        self.assertEqual(a.last_name, 'Villanueva')
        self.assertEqual(a.first_name, 'Nena')
        self.assertEqual(a.middle_name, '')
        self.assertEqual(a.middle_initial, '')

    def test_four_part_name_keeps_both_middle_tokens(self):
        a = self._app('Juan Miguel Santos Cruz')
        self.assertEqual(a.last_name, 'Cruz')
        self.assertEqual(a.first_name, 'Juan')
        self.assertEqual(a.middle_name, 'Miguel Santos')

    def test_a_single_token_is_treated_as_the_surname(self):
        a = self._app('Madonna')
        self.assertEqual(a.last_name, 'Madonna')
        self.assertEqual(a.first_name, '')

    def test_an_empty_name_does_not_raise(self):
        a = self._app('')
        self.assertEqual((a.last_name, a.first_name, a.middle_name), ('', '', ''))
