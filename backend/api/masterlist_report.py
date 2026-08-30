"""VPSEA scholars masterlist — the BiPSU "LIST OF SCHOLARS" Word document.

The office's own document is bundled at ``templates/docx/masterlist_template.docx``
as a docxtpl (Jinja) template: every table body row carries ``{%tr for row in
programN.female %}`` loops and ``{{ row.* }}`` placeholders. Rendering fills the
rows in place, so the letterhead, column headers, merged cells, borders, page
setup and signatory blocks are the office's originals rather than anything
recreated in code.

The template exposes 16 program slots. ``PROGRAM_SLOTS`` below decides which
scholarship lands in which slot — reorder that list to reorder the document.
"""
import os
import re

TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'templates', 'docx', 'masterlist_template.docx',
)

# Slot -> (heading, scholarship type, layout).
#   'gendered' fills programN.female / programN.male
#   'students' fills programN.students (the BiPSU Staff table has no gender split)
# program11 is skipped: its table only has a female loop, so a program placed
# there would silently drop its male scholars.
# (slot, heading, source key, layout, header style)
PROGRAM_SLOTS = [
    ('program1',  'ACADEMIC',        'Academic',    'gendered', 'academic'),
    ('program2',  'BiPSU STAFF',     'Staff',       'students', 'staff'),
    ('program3',  'AFFIRMATIVE',     'Affirmative', 'gendered', 'award'),
    ('program4',  'CHED FULL MERIT', 'CHED_FULL',   'gendered', 'award'),
    ('program5',  'CHED HALF MERIT', 'CHED_HALF',   'gendered', 'award'),
    ('program6',  'DOST',            'DOST',        'gendered', 'award'),
    ('program7',  'TDP',             'TDP',         'gendered', 'award'),
    ('program8',  'TES',             'TES',         'gendered', 'award'),
    ('program9',  'GSIS',            'GSIS',        'gendered', 'award'),
    ('program10', 'CoScho',          'CoScho',      'gendered', 'award'),
    ('program12', 'SPORTS',          'Sports',      'gendered', 'award'),
]
ALL_SLOTS = [f'program{i}' for i in range(1, 17)]

# The office template carries 16 program blocks but PROGRAM_SLOTS fills fewer.
# Spare slots are named with this marker so the rendered document can find and
# delete them — otherwise they print as an unnamed heading over an empty table.
UNUSED_MARKER = '∅UNUSED'

# Every column heading in the office template maps to one row field. The
# preview builds its cells from this, so a table's columns always match the
# document's — no slot's headings are hardcoded here.
HEADER_FIELD = {
    'NO.': 'no',
    'AWARD NUMBER': 'award_number',
    'LAST NAME': 'last_name',
    'FIRST NAME': 'first_name',
    'MIDDLE NAME': 'middle_name',
    'M.I.': 'm_i',
    'SEX': 'sex',
    'BRGY./ST.': 'brgy_st',
    'MUN.': 'mun',
    'MUNICIPALITY': 'municipality',
    'PROV.': 'prov',
    'PROVINCE': 'province',
    'CONG. DIST.': 'cong_dist',
    'COURSE': 'course',
    'YR.': 'yr',
    'YEAR LEVEL': 'year_level',
    'GWA': 'gwa',
    '%': 'percent',
    'NUMBER': 'number',
    'SCHOLARSHIP PROGRAM': 'scholarship_program',
    'PROGRAM': 'program',
}

_HEADER_CACHE = {}


def slot_headers():
    """The real column headings of each program table, read from the template.

    Every slot has its own shape — GSIS carries no award number, Sports has
    neither barangay nor congressional district — so reading them beats keeping
    a second copy in code that can drift out of step with the document.
    """
    if _HEADER_CACHE:
        return _HEADER_CACHE
    import re
    import docx

    tag = re.compile(r'\{%tr\s*for row in (\w+)\.(\w+)')
    document = docx.Document(TEMPLATE_PATH)
    for table in document.tables:
        joined = ' '.join(c.text for r in table.rows for c in r.cells)
        found = tag.search(joined)
        if not found or found.group(1) in _HEADER_CACHE:
            continue
        raw = [c.text.strip().replace(chr(10), ' ') for c in table.rows[1].cells]
        headings = []
        for h in raw:                      # collapse Word's merged duplicates
            if not headings or headings[-1] != h:
                headings.append(h)
        _HEADER_CACHE[found.group(1)] = headings
    return _HEADER_CACHE


def cells_for(row, headings):
    """A row rendered in the order a given table's headings ask for."""
    return [row.get(HEADER_FIELD.get(h, ''), '') for h in headings]


FEMALE_VALUES = ('F', 'FEMALE')


def _is_female(gender):
    return (gender or '').strip().upper() in FEMALE_VALUES


def _initial(name):
    name = (name or '').strip()
    return f'{name[0].upper()}.' if name else ''


def _blank_row(no):
    """Every field any of the 16 tables might reference, so a row dropped into
    any slot fills the cells that exist there and leaves the rest blank."""
    return {
        'no': no, 'col': '',
        'last_name': '', 'first_name': '', 'first_name1': '',
        'middle_name': '', 'm_i': '',
        'sex': '', 'brgy_st': '', 'mun': '', 'municipality': '',
        'prov': '', 'province': '',
        'course': '', 'yr': '', 'year_level': '',
        'gwa': '', 'percent': '', 'number': '',
        'award_number': '', 'cong_dist': '',
        'scholarship_program': '', 'program': '',
    }


def _application_row(no, app):
    """A row from an approved Application (the student-portal programs)."""
    p = app.student
    u = p.user
    middle = getattr(p, 'middle_name', '') or ''
    gwa = p.gwa or 0
    if app.scholarship.type == 'Academic':
        percent = 'Univ. Scholar' if gwa <= 1.29 else ('College Scholar' if gwa <= 1.50 else '')
    else:
        percent = ''

    row = _blank_row(no)
    row.update({
        'last_name': u.last_name or '',
        'first_name': u.first_name or '',
        'first_name1': u.first_name or '',
        'middle_name': middle,
        'm_i': _initial(middle),
        'sex': (p.gender or '')[:1].upper(),
        'brgy_st': p.barangay or '',
        'mun': p.municipality or '',
        'municipality': p.municipality or '',
        'prov': p.province or '',
        'province': p.province or '',
        'course': p.course or '',
        'yr': p.year_level or '',
        'year_level': p.year_level or '',
        'gwa': f'{gwa:.2f}' if gwa else '',
        'percent': percent,
        'number': p.student_id or '',
        'award_number': app.award_number,
        'cong_dist': app.congress_district,
        'scholarship_program': app.scholarship.name,
        'program': app.scholarship.name,
    })
    return row


def _affirmative_row(no, app):
    """A row from an approved AffirmativeStaffApplication (Affirmative / Staff)."""
    parts = (app.full_name or '').strip().split()
    if len(parts) >= 3:
        last, first, middle = parts[-1], parts[0], parts[1]
    elif len(parts) == 2:
        last, first, middle = parts[-1], parts[0], ''
    else:
        last, first, middle = (parts[0] if parts else ''), '', ''

    name = 'BiPSU Staff Scholarship' if app.is_nsu_staff else 'Affirmative Action Scholarship'
    row = _blank_row(no)
    row.update({
        'last_name': last, 'first_name': first, 'first_name1': first,
        'middle_name': middle, 'm_i': _initial(middle),
        'sex': (app.gender or '')[:1].upper(),
        'brgy_st': app.barangay or '',
        'mun': app.municipality or '', 'municipality': app.municipality or '',
        'prov': app.province or '', 'province': app.province or '',
        'course': app.course or '',
        'yr': app.year_level or '', 'year_level': app.year_level or '',
        'percent': '100%' if app.is_nsu_staff else '75%',
        'number': app.student_id or '',
        'scholarship_program': name, 'program': name,
    })
    return row


def _imported_row(no, rec):
    """A row from an ImportedScholar — a scholar the office uploaded or typed in.

    These are the bulk of the office's records for every programme without a
    portal, and the masterlist used to read straight past them: it asked only
    Application and AffirmativeStaffApplication, so a list of imported CHED or
    DOST scholars produced an empty table under a printed heading.
    """
    middle = rec.middle_name or ''
    gwa = rec.gwa or 0
    if rec.scholarship_type == 'Academic':
        percent = 'Univ. Scholar' if 0 < gwa <= 1.29 else ('College Scholar' if 0 < gwa <= 1.50 else '')
    else:
        percent = ''

    # Set in _sources so the programme's proper name costs one query, not one
    # per scholar. Falls back to the bare type if a Scholarship row is missing.
    name = getattr(rec, '_program_name', '') or rec.scholarship_type

    row = _blank_row(no)
    row.update({
        'last_name': rec.last_name or '',
        'first_name': rec.first_name or '',
        'first_name1': rec.first_name or '',
        'middle_name': middle,
        'm_i': _initial(middle),
        'sex': (rec.gender or '')[:1].upper(),
        'brgy_st': rec.barangay or '',
        'mun': rec.municipality or '', 'municipality': rec.municipality or '',
        'prov': rec.province or '', 'province': rec.province or '',
        'course': rec.course or '',
        'yr': rec.year_level or '', 'year_level': rec.year_level or '',
        'gwa': f'{gwa:.2f}' if gwa else '',
        'percent': percent,
        'number': rec.student_id or '',
        'award_number': rec.award_number or '',
        'cong_dist': rec.congress_district or '',
        'scholarship_program': name, 'program': name,
    })
    return row


def _row_for(no, record):
    """Build a row from whichever of the three shapes a scholar arrived in."""
    from .models import AffirmativeStaffApplication, ImportedScholar

    if isinstance(record, AffirmativeStaffApplication):
        return _affirmative_row(no, record)
    if isinstance(record, ImportedScholar):
        return _imported_row(no, record)
    return _application_row(no, record)


def _gender_of(record):
    """The gender to split a table on, whatever shape the record is."""
    from .models import Application

    if isinstance(record, Application):
        return record.student.gender
    return record.gender


def _sources(term_label=None):
    """Every scholarship type mapped to its approved records, already ordered.

    Three shapes land here. Applications are the portal's awards;
    AffirmativeStaffApplications cover the two programmes applied for outside it;
    ImportedScholars are everyone the office uploaded or added by hand, which for
    most programmes is nearly all of them.

    Imported rows are scoped to one term because they carry one, and a document
    stamped with a single semester should not print every scholar the office has
    ever imported. A row already claimed by a student account is skipped — that
    scholar has an Application above and would otherwise be printed twice.
    """
    from .models import (Application, AffirmativeStaffApplication, ImportedScholar,
                         Scholarship, SystemSettings, split_ched)

    if term_label is None:
        settings_obj, _ = SystemSettings.objects.get_or_create(pk=1)
        term_label = settings_obj.academic_year

    programme_names = dict(Scholarship.objects.values_list('type', 'name'))

    def imported(stype):
        rows = list(
            ImportedScholar.objects
            .filter(scholarship_type=stype, term_label=term_label, claimed_by__isnull=True)
            .order_by('last_name', 'first_name')
        )
        for r in rows:
            r._program_name = programme_names.get(stype, stype)
        return rows

    def apps(stype):
        return list(
            Application.objects.filter(status='Approved', scholarship__type=stype)
            .select_related('student__user', 'scholarship')
            .order_by('student__user__last_name', 'student__user__first_name')
        ) + imported(stype)

    ched_full, ched_half = split_ched(apps('CHED'))

    def affirmative(qualified_for):
        return list(
            AffirmativeStaffApplication.objects.filter(
                status='Approved', qualified_for=qualified_for
            ).order_by('full_name')
        ) + imported(qualified_for)

    return {
        'Academic': apps('Academic'),
        'DOST': apps('DOST'),
        'GSIS': apps('GSIS'),
        'TDP': apps('TDP'),
        'TES': apps('TES'),
        'CoScho': apps('CoScho'),
        'Sports': apps('Sports'),
        'CHED_FULL': ched_full,
        'CHED_HALF': ched_half,
        'Affirmative': affirmative('Affirmative'),
        'Staff': affirmative('Staff'),
    }


def build_context(sources=None):
    """The docxtpl context: one entry per program slot, plus per-slot counts."""
    sources = sources if sources is not None else _sources()
    headings = slot_headers()
    context = {slot: {'name': UNUSED_MARKER, 'female': [], 'male': [], 'students': []}
               for slot in ALL_SLOTS}
    summary = []

    for slot, heading, key, layout, header_style in PROGRAM_SLOTS:
        records = sources.get(key, [])
        # A slot now holds a mix of shapes — an imported row can sit beside a
        # portal application in the same table — so the builder is chosen per
        # record rather than per slot.
        build, gender_of = _row_for, _gender_of

        entry = {'name': heading, 'female': [], 'male': [], 'students': []}
        if layout == 'students':
            entry['students'] = [build(i, r) for i, r in enumerate(records, 1)]
        else:
            female = [r for r in records if _is_female(gender_of(r))]
            male = [r for r in records if not _is_female(gender_of(r))]
            entry['female'] = [build(i, r) for i, r in enumerate(female, 1)]
            entry['male'] = [build(i, r) for i, r in enumerate(male, 1)]

        context[slot] = entry
        summary.append({
            'slot': slot, 'heading': heading, 'key': key, 'layout': layout,
            'headers': headings.get(slot, []),
            'female': len(entry['female']), 'male': len(entry['male']),
            'students': len(entry['students']),
            'total': len(records),
        })

    return context, summary


def _restamp_period(document, sy, semester):
    """The template's headings carry a literal '1st SEMESTER SY: 2019-2020'.

    There is no Jinja tag for it, so rewrite those headings after rendering
    rather than shipping a document stamped with the wrong school year.
    """
    second = semester.strip().lower().startswith('2')
    year_re = re.compile(r'SY:\s*\d{4}\s*[-–]\s*\d{4}', re.I)
    sem_re = re.compile(r'\b1(st|ST)\b')

    def fix(text):
        text = year_re.sub(f'SY: {sy}', text)
        if second:
            text = sem_re.sub(lambda m: '2ND' if m.group(1).isupper() else '2nd', text)
        return text

    def walk(paragraphs):
        for p in paragraphs:
            original = p.text
            if 'SY:' not in original.upper():
                continue
            updated = fix(original)
            if updated == original or not p.runs:
                continue
            # Collapse into the first run so a match spanning runs still lands.
            p.runs[0].text = updated
            for run in p.runs[1:]:
                run.text = ''

    walk(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                walk(cell.paragraphs)
    for section in document.sections:
        walk(section.header.paragraphs)
        walk(section.footer.paragraphs)


def _drop_unused_sections(document):
    """Delete the spare program blocks the office template ships with.

    A program block runs from its heading paragraph to the next heading. A
    heading is identified structurally: it is the paragraph immediately followed
    by the template's "SCHOLARSHIP GRANT" line.
    """
    from docx.text.paragraph import Paragraph

    body = document.element.body
    children = [c for c in body.iterchildren()
                if c.tag.endswith('}p') or c.tag.endswith('}tbl')]

    def is_para(el):
        return el.tag.endswith('}p')

    def text_of(el):
        return Paragraph(el, document).text.strip() if is_para(el) else ''

    headings = []
    for i, el in enumerate(children):
        if not is_para(el):
            continue
        for nxt in children[i + 1:]:
            if not is_para(nxt):
                break
            following = text_of(nxt)
            if not following:
                continue
            if following.upper().startswith('SCHOLARSHIP GRANT'):
                headings.append(i)
            break

    removed = 0
    for pos, start in enumerate(headings):
        if UNUSED_MARKER not in text_of(children[start]):
            continue
        end = headings[pos + 1] if pos + 1 < len(headings) else len(children)
        for el in children[start:end]:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
                removed += 1

    # Belt and braces: never let the marker itself reach the reader.
    for el in children:
        if is_para(el) and UNUSED_MARKER in text_of(el):
            para = Paragraph(el, document)
            for run in para.runs:
                run.text = run.text.replace(UNUSED_MARKER, '')
    return removed


def build_document(academic_year, semester):
    """Return ``(BytesIO, summary)`` for the filled masterlist document."""
    from io import BytesIO
    from docxtpl import DocxTemplate

    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f'The masterlist template is missing at {TEMPLATE_PATH}. '
            'Restore it from the office copy before generating this report.'
        )

    context, summary = build_context()
    tpl = DocxTemplate(TEMPLATE_PATH)
    tpl.render(context)
    _drop_unused_sections(tpl.docx)
    _restamp_period(tpl.docx, academic_year, semester)

    buf = BytesIO()
    tpl.save(buf)
    buf.seek(0)
    return buf, summary
